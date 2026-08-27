from docx import Document


doc = Document("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")

capture = False
imgs = 0
captions = []
heads = []

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    style = paragraph.style.name if paragraph.style else ""
    if text == "5.4 Diagramas de código":
        capture = True
    elif capture and text == "Capítulo 6. Desarrollo":
        heads.append((i, style, text))
        break

    if capture:
        if paragraph._p.xpath(".//a:blip"):
            imgs += len(paragraph._p.xpath(".//a:blip"))
        if style == "Caption" and text:
            captions.append(text)
        if (
            text.startswith("5.4")
            or text.startswith("Tabla “")
            or text.startswith("Consideraciones")
            or style.startswith("Título")
            or style == "Heading 4"
        ):
            heads.append((i, style, text))

body = list(doc.element.body)
start_el = end_el = None
for paragraph in doc.paragraphs:
    if paragraph.text.strip() == "5.4 Diagramas de código":
        start_el = paragraph._p
    if paragraph.text.strip() == "Capítulo 6. Desarrollo":
        end_el = paragraph._p
        break

tables = 0
in_range = False
for element in body:
    if element is start_el:
        in_range = True
        continue
    if element is end_el:
        break
    if in_range and element.tag.endswith("}tbl"):
        tables += 1

print("HEADINGS")
for item in heads:
    print(item)
print("IMAGES", imgs)
print("CAPTIONS", len(captions))
for caption in captions:
    print("CAP", caption)
print("TABLES_IN_54", tables)
