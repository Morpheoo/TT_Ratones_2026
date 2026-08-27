from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")
IMG_DIR = Path("reportes/figuras/secuencia_53")


INTRO = [
    (
        "Los diagramas de secuencia describen la interacción temporal entre los actores "
        "del prototipo, la interfaz Streamlit, los servicios internos y la base de datos "
        "PostgreSQL para cada caso de uso definido en la sección 5.2. A diferencia de los "
        "diagramas de casos de uso, que muestran qué funcionalidades existen y quiénes "
        "participan en ellas, los diagramas de secuencia detallan el orden de los mensajes "
        "necesarios para ejecutar cada flujo principal y sus trayectorias alternativas relevantes."
    ),
    (
        "Para mantener la trazabilidad con los requerimientos funcionales, esta sección presenta "
        "un diagrama por caso de uso, desde CU1 hasta CU10. En todos los diagramas se conserva "
        "una estructura homogénea: actor, interfaz o módulo visible, servicio de aplicación, "
        "componente técnico especializado cuando aplica, y persistencia en PostgreSQL o archivos auxiliares."
    ),
]


FIGURES = [
    ("fig_seq_cu01_registrar_usuario.png", "Diagrama de secuencia CU1: Registrar usuario"),
    ("fig_seq_cu02_iniciar_sesion.png", "Diagrama de secuencia CU2: Iniciar sesión"),
    ("fig_seq_cu03_gestionar_tratamientos.png", "Diagrama de secuencia CU3: Gestionar tratamientos"),
    ("fig_seq_cu04_registrar_experimento.png", "Diagrama de secuencia CU4: Registrar experimento"),
    ("fig_seq_cu05_configurar_rois.png", "Diagrama de secuencia CU5: Configurar ROIs"),
    ("fig_seq_cu06_ejecutar_analisis.png", "Diagrama de secuencia CU6: Ejecutar análisis de video"),
    ("fig_seq_cu07_agrupar_comportamientos.png", "Diagrama de secuencia CU7: Agrupar comportamientos"),
    ("fig_seq_cu08_editar_tiempos.png", "Diagrama de secuencia CU8: Editar tiempos conductuales"),
    ("fig_seq_cu09_exportar_resultados.png", "Diagrama de secuencia CU9: Exportar resultados"),
    ("fig_seq_cu10_bitacora_auditoria.png", "Diagrama de secuencia CU10: Consultar bitácora de auditoría"),
]


def paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_run_font(run, size_pt=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def set_paragraph_tnr(paragraph, size_pt=12):
    for run in paragraph.runs:
        set_run_font(run, size_pt=size_pt)


def main():
    doc = Document(DOCX)

    start = None
    end = None
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "5.3 Diagramas de secuencia":
            start = i
        elif start is not None and text == "5.4 Diagramas de código":
            end = i
            break

    if start is None or end is None:
        raise RuntimeError("No se encontró el rango 5.3 -> 5.4.")

    heading = doc.paragraphs[start]

    # Delete old 5.3 body while preserving the 5.3 heading and 5.4 heading.
    for paragraph in list(doc.paragraphs[start + 1:end]):
        delete_paragraph(paragraph)

    anchor = heading

    for text in INTRO:
        p = paragraph_after(anchor, style="Normal")
        run = p.add_run(text)
        set_run_font(run, 12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        anchor = p

    for file_name, caption in FIGURES:
        img_path = IMG_DIR / file_name
        if not img_path.exists():
            raise FileNotFoundError(img_path)

        img_p = paragraph_after(anchor, style="Normal")
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_p.add_run()
        img_run.add_picture(str(img_path), width=Inches(6.35))
        anchor = img_p

        cap_p = paragraph_after(anchor, style="Caption")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption + " [autoría propia]")
        set_run_font(cap_run, 10)
        anchor = cap_p

    doc.save(DOCX)
    print("Sección 5.3 reemplazada correctamente.")


if __name__ == "__main__":
    main()
