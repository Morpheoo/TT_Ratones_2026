from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")

BLOCKS = [
    (
        "7.4.4 Modelos y estrategias experimentales exploradas",
        [
            (
                "Además del clasificador Random Forest productivo, se exploraron estrategias "
                "complementarias para reducir los fallos de generalización observados en "
                "Grooming. B-SOiD se evaluó como alternativa no supervisada basada en motivos "
                "conductuales, ya que permite agrupar patrones de movimiento sin depender "
                "exclusivamente de etiquetas frame a frame. En la validación Leave-One-Out "
                "con N = 26, B-SOiD obtuvo un F1-Score justo de 0.357 para Grooming, inferior "
                "al Random Forest de SimBA (0.400), pero mostró capacidad de rescate en videos "
                "donde SimBA colapsaba y predecía muy pocos frames positivos."
            ),
            (
                "A partir de este comportamiento se evaluó Ensemble OR, que considera positiva "
                "una predicción si SimBA o B-SOiD detectan la conducta. Esta estrategia elevó "
                "el F1-Score justo de Grooming a 0.502, pero también incrementó falsos positivos "
                "en videos donde SimBA ya funcionaba correctamente. Por ello se diseñó la "
                "estrategia Conditional, que mantiene SimBA como clasificador principal y activa "
                "B-SOiD únicamente cuando SimBA predice menos de 250 frames positivos. Esta regla "
                "obtuvo el mejor desempeño para Grooming, con F1-Score justo de 0.523, y representa "
                "el equilibrio más favorable entre rescate de eventos y control de falsos positivos."
            ),
            (
                "También se evaluó una estrategia Dynamic, basada en ajustar el umbral de Grooming "
                "por video cuando SimBA predice pocos positivos. En la validación global alcanzó "
                "un F1-Score justo de 0.507, cercano a Conditional, pero su análisis en videos "
                "críticos mostró que bajar el umbral no siempre recupera la conducta cuando las "
                "características no separan adecuadamente Grooming del fondo. Por esta razón se "
                "mantuvo como herramienta diagnóstica y no como regla productiva principal."
            ),
            (
                "La capa LSTM se exploró como mecanismo de rescate temporal para eventos breves "
                "o fragmentados. A diferencia del Random Forest, que clasifica frame a frame sobre "
                "características tabulares, la LSTM modela continuidad temporal en ventanas de "
                "frames. Su función en el sistema no es reemplazar SimBA, sino elevar predicciones "
                "cuando el Random Forest queda en una zona de incertidumbre. En las pruebas internas "
                "con cinco videos críticos, la LSTM reentrenada con 26 videos alcanzó F1 promedio "
                "de 0.959 con umbral 0.75, 0.946 con umbral 0.50 y 0.885 con umbral 0.11; estos "
                "valores no se reportan como validación ciega, sino como evidencia de que el backend "
                "temporal reconoce los patrones de Grooming cuando está correctamente integrado."
            ),
            (
                "Finalmente, se probaron estrategias adicionales como mirror augmentation y bagging "
                "multi-semilla. Mirror augmentation duplicó el conjunto efectivo a 52 archivos "
                "(26 originales y 26 espejados), pero no mejoró de forma consistente la validación "
                "ciega en videos críticos. Bagging multi-semilla tampoco corrigió los colapsos de "
                "SimBA: en cinco videos críticos obtuvo F1 promedio de 0.010 frente a 0.017 del "
                "baseline SimBA en el mismo subconjunto. Este resultado sugiere que el problema no "
                "era únicamente varianza por semilla, sino falta de separabilidad para ciertos estilos "
                "de Grooming dentro del espacio de características disponible."
            ),
        ],
    ),
    (
        "7.4.5 Decisión técnica final",
        [
            (
                "Con base en estos resultados, la arquitectura final conserva SimBA Random Forest "
                "como núcleo productivo por su estabilidad, interpretabilidad y compatibilidad con "
                "datos tabulares; utiliza Conditional como mejor estrategia para Grooming cuando "
                "SimBA pierde sensibilidad; mantiene SimBA RF como método principal para Thigmotaxis, "
                "donde obtuvo F1-Score justo de 0.636; y conserva LSTM rescue como apoyo temporal "
                "para eventos fragmentados. Las estrategias Dynamic, mirror augmentation y bagging "
                "se documentan como experimentos útiles, pero no sustituyen al flujo productivo por "
                "no superar de forma consistente a Conditional en validación ciega."
            ),
            (
                "Esta decisión técnica evita presentar el sistema como un único clasificador perfecto "
                "y, en cambio, documenta una arquitectura evaluada de forma comparativa. El resultado "
                "defendible es que el prototipo identifica qué componente funciona mejor para cada "
                "conducta: SimBA Random Forest para Thigmotaxis, Conditional para Grooming y LSTM "
                "como rescate temporal complementario. Así, la elección final responde a evidencia "
                "experimental y no a preferencia arbitraria por un algoritmo."
            ),
        ],
    ),
]


def set_tnr(run, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_after(anchor, text, style):
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    p = Paragraph(new_p, anchor._parent)
    p.style = style
    run = p.add_run(text)
    set_tnr(run)
    return p


def find(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p.text.strip()):
            return p
    return None


def delete_between(start, end):
    cur = start._element.getnext()
    end_el = end._element
    while cur is not None and cur is not end_el:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt


def main():
    doc = Document(DOCX)
    last_original = find(doc, lambda t: t.startswith("Finalmente, la limitación es remediable"))
    conclusions = find(doc, lambda t: t == "Conclusiones")
    if last_original is None or conclusions is None:
        raise RuntimeError("No se encontró el bloque de referencia para corregir 7.4.")

    delete_between(last_original, conclusions)
    anchor = last_original
    for heading, paragraphs in BLOCKS:
        anchor = add_after(anchor, heading, "Título 3 TNR")
        for text in paragraphs:
            anchor = add_after(anchor, text, "Texto Normal TNR")

    doc.save(DOCX)
    print("Orden de 7.4.4 y 7.4.5 corregido.")


if __name__ == "__main__":
    main()
