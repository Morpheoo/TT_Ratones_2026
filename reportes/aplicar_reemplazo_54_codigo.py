from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")
IMG_DIR = Path("reportes/figuras/codigo_54")


INTRO_54 = [
    (
        "En esta sección se documenta la estructura interna del prototipo a partir "
        "de los módulos de código y del modelo de datos implementado. Dado que la "
        "aplicación fue desarrollada en Python con Streamlit, varias responsabilidades "
        "del sistema se organizan como módulos funcionales y páginas de aplicación, "
        "no como clases orientadas a objetos tradicionales. Por ello, el diagrama de "
        "código representa componentes reales del repositorio, sus responsabilidades "
        "principales y sus dependencias con la capa de persistencia."
    )
]

TEXT_541 = [
    (
        "El diagrama de módulos resume la relación entre la interfaz Streamlit, los "
        "servicios internos y la persistencia. La capa de interfaz se compone de "
        "Home.py y de las páginas ubicadas en pages/, entre ellas el inicio de sesión, "
        "la ingesta de video, la configuración de zonas, el análisis final, los "
        "resultados y el panel administrativo."
    ),
    (
        "La lógica de aplicación se concentra en módulos verificables del repositorio: "
        "src/auth.py gestiona registro, autenticación, OTP, recuperación de contraseña "
        "y roles; src/treatments.py administra el catálogo de tratamientos; "
        "src/db/experiment_history.py registra experimentos y configuraciones ROI; "
        "src/simba_roi_bridge.py sincroniza las regiones del laberinto con SimBA; "
        "src/db/behavior_edits.py registra las ediciones manuales de tiempos "
        "conductuales; y src/security_logger.py persiste eventos de seguridad."
    ),
    (
        "La persistencia se realiza mediante PostgreSQL, inicializado con schema.sql "
        "y accedido desde src/db/connection.py. Además de la base de datos, el sistema "
        "utiliza archivos auxiliares para videos, salidas del pipeline y configuraciones "
        "de zonas, como zonas_activas.json."
    ),
]

TEXT_542 = [
    (
        "El modelo entidad-relación documenta las siete tablas activas del prototipo. "
        "La tabla users almacena credenciales, rol, estado de verificación, estado "
        "activo y datos de perfil. La tabla treatments contiene el catálogo de "
        "tratamientos experimentales. La tabla experiments registra cada sesión EPM "
        "con su identificador de roedor, tratamiento, responsable, ruta de video y "
        "estado de procesamiento."
    ),
    (
        "Las regiones de interés se almacenan en roi_configurations mediante una "
        "relación uno a muchos con experiments. Los resultados generados por el "
        "pipeline se guardan en analysis_results, incluyendo distancia total, tiempos "
        "en brazos abiertos, brazos cerrados y centro, así como duración de Grooming "
        "y Thigmotaxis."
    ),
    (
        "La trazabilidad se cubre con dos tablas complementarias. security_audit_log "
        "registra eventos de seguridad, autenticación y verificación. behavior_edits "
        "almacena snapshots before/after de las ediciones manuales de tiempos "
        "conductuales, asociados al experimento y al usuario que realizó la corrección."
    ),
]

DICTIONARY = [
    (
        "users",
        "Tabla de usuarios del sistema. Almacena credenciales, rol, verificación, estado de cuenta y datos de perfil institucional.",
        [
            ("id", "SERIAL", "PRIMARY KEY", "Auto-incremental", "Identificador único del usuario."),
            ("username", "VARCHAR(50)", "UNIQUE, NOT NULL", "Sin valor", "Correo institucional usado como nombre de usuario."),
            ("password_hash", "TEXT", "NOT NULL", "Sin valor", "Hash bcrypt de la contraseña."),
            ("role", "VARCHAR(20)", "CHECK", "investigador", "Rol operativo: admin, investigador o estudiante."),
            ("is_verified", "BOOLEAN", "-", "FALSE", "Indica si la cuenta completó verificación OTP o fue verificada como admin."),
            ("verification_code", "VARCHAR(6)", "-", "NULL", "Código OTP temporal."),
            ("verification_code_created_at", "TIMESTAMP", "-", "NULL", "Fecha y hora de emisión del OTP."),
            ("is_active", "BOOLEAN", "-", "TRUE", "Permite suspender o habilitar la cuenta."),
            ("created_at", "TIMESTAMP", "-", "CURRENT_TIMESTAMP", "Fecha de creación de la cuenta."),
            ("full_name", "VARCHAR(200)", "-", "NULL", "Nombre completo del usuario."),
            ("accepted_terms", "BOOLEAN", "-", "FALSE", "Aceptación de términos y condiciones."),
            ("boleta", "VARCHAR(20)", "-", "NULL", "Identificador institucional de estudiante."),
            ("carrera", "VARCHAR(150)", "-", "NULL", "Programa académico del estudiante."),
            ("escuela", "VARCHAR(100)", "-", "NULL", "Unidad académica del estudiante."),
            ("num_empleado", "VARCHAR(20)", "-", "NULL", "Número de empleado del investigador/docente."),
            ("area", "VARCHAR(150)", "-", "NULL", "Área académica o de investigación."),
            ("centro", "VARCHAR(100)", "-", "NULL", "Centro o unidad de adscripción."),
        ],
    ),
    (
        "treatments",
        "Catálogo de tratamientos experimentales disponibles para registrar experimentos.",
        [
            ("id", "SERIAL", "PRIMARY KEY", "Auto-incremental", "Identificador único del tratamiento."),
            ("name", "VARCHAR(100)", "UNIQUE, NOT NULL", "Sin valor", "Nombre del tratamiento."),
            ("description", "TEXT", "-", "NULL", "Descripción opcional del tratamiento."),
            ("created_by", "INTEGER", "FK users(id)", "NULL", "Usuario que creó el tratamiento."),
            ("created_at", "TIMESTAMP", "-", "CURRENT_TIMESTAMP", "Fecha de creación."),
            ("is_active", "BOOLEAN", "-", "TRUE", "Estado lógico del tratamiento."),
        ],
    ),
    (
        "experiments",
        "Registro principal de sesiones experimentales EPM y videos asociados.",
        [
            ("id", "SERIAL", "PRIMARY KEY", "Auto-incremental", "Identificador único del experimento."),
            ("rat_id", "VARCHAR(50)", "NOT NULL", "Sin valor", "Identificador del roedor."),
            ("treatment", "VARCHAR(50)", "NOT NULL", "Sin valor", "Tratamiento aplicado."),
            ("experiment_date", "DATE", "-", "NULL", "Fecha de la sesión experimental."),
            ("responsible", "VARCHAR(100)", "-", "NULL", "Responsable del experimento."),
            ("video_path", "TEXT", "NOT NULL", "Sin valor", "Ruta del video experimental."),
            ("duration_seconds", "FLOAT", "-", "NULL", "Duración del video en segundos."),
            ("created_by", "INTEGER", "FK users(id)", "NULL", "Usuario que registró el experimento."),
            ("processed", "BOOLEAN", "-", "FALSE", "Indica si el experimento fue procesado."),
            ("created_at", "TIMESTAMP", "-", "CURRENT_TIMESTAMP", "Fecha de registro."),
        ],
    ),
    (
        "roi_configurations",
        "Configuración de regiones de interés del laberinto EPM asociadas a un experimento.",
        [
            ("id", "SERIAL", "PRIMARY KEY", "Auto-incremental", "Identificador único de la configuración."),
            ("experiment_id", "INTEGER", "FK experiments(id), ON DELETE CASCADE", "NULL", "Experimento asociado."),
            ("zone_type", "VARCHAR(50)", "NOT NULL", "Sin valor", "Tipo o nombre de zona EPM."),
            ("coordinates_json", "JSONB", "NOT NULL", "Sin valor", "Coordenadas de la ROI en formato JSON."),
            ("scale_factor", "FLOAT", "-", "NULL", "Factor de escala respecto al video original."),
        ],
    ),
    (
        "analysis_results",
        "Resultados cuantitativos generados por el pipeline de análisis conductual.",
        [
            ("id", "SERIAL", "PRIMARY KEY", "Auto-incremental", "Identificador único del resultado."),
            ("experiment_id", "INTEGER", "FK experiments(id), ON DELETE CASCADE", "NULL", "Experimento analizado."),
            ("timestamp", "TIMESTAMP", "-", "CURRENT_TIMESTAMP", "Fecha de generación del resultado."),
            ("total_distance", "FLOAT", "-", "0.0", "Distancia total recorrida."),
            ("time_open_arms", "FLOAT", "-", "0.0", "Tiempo en brazos abiertos."),
            ("time_closed_arms", "FLOAT", "-", "0.0", "Tiempo en brazos cerrados."),
            ("time_center", "FLOAT", "-", "0.0", "Tiempo en zona central."),
            ("head_dips_count", "INTEGER", "-", "0", "Conteo de head dips."),
            ("rearing_count", "INTEGER", "-", "0", "Conteo de rearing."),
            ("grooming_duration", "FLOAT", "-", "0.0", "Duración acumulada de Grooming."),
            ("thigmotaxis_duration", "FLOAT", "-", "0.0", "Duración acumulada de Thigmotaxis."),
            ("status", "VARCHAR(20)", "-", "pending", "Estado del resultado."),
            ("trajectory_path", "TEXT", "Migración", "NULL", "Ruta del archivo de trayectoria generado por la aplicación."),
        ],
    ),
    (
        "behavior_edits",
        "Auditoría de ediciones manuales de tiempos conductuales con snapshot before/after.",
        [
            ("id", "SERIAL", "PRIMARY KEY", "Auto-incremental", "Identificador único de la edición."),
            ("experiment_id", "INTEGER", "FK experiments(id), ON DELETE CASCADE", "Sin valor", "Experimento corregido."),
            ("edited_by", "INTEGER", "FK users(id), ON DELETE SET NULL", "NULL", "Usuario que realizó la edición."),
            ("edited_by_email", "TEXT", "-", "NULL", "Correo del usuario editor."),
            ("edited_role", "TEXT", "-", "NULL", "Rol del usuario editor."),
            ("edited_at", "TIMESTAMP", "NOT NULL", "CURRENT_TIMESTAMP", "Fecha de la edición."),
            ("before_open", "FLOAT", "-", "NULL", "Tiempo previo en brazos abiertos."),
            ("before_closed", "FLOAT", "-", "NULL", "Tiempo previo en brazos cerrados."),
            ("before_center", "FLOAT", "-", "NULL", "Tiempo previo en centro."),
            ("before_grooming", "FLOAT", "-", "NULL", "Tiempo previo de Grooming."),
            ("before_thigmo", "FLOAT", "-", "NULL", "Tiempo previo de Thigmotaxis."),
            ("after_open", "FLOAT", "-", "NULL", "Tiempo corregido en brazos abiertos."),
            ("after_closed", "FLOAT", "-", "NULL", "Tiempo corregido en brazos cerrados."),
            ("after_center", "FLOAT", "-", "NULL", "Tiempo corregido en centro."),
            ("after_grooming", "FLOAT", "-", "NULL", "Tiempo corregido de Grooming."),
            ("after_thigmo", "FLOAT", "-", "NULL", "Tiempo corregido de Thigmotaxis."),
            ("note", "TEXT", "-", "NULL", "Justificación de la edición."),
        ],
    ),
    (
        "security_audit_log",
        "Bitácora de eventos de seguridad, autenticación y verificación del sistema.",
        [
            ("id", "SERIAL", "PRIMARY KEY", "Auto-incremental", "Identificador único del evento."),
            ("timestamp", "TIMESTAMP", "NOT NULL", "CURRENT_TIMESTAMP", "Fecha y hora del evento."),
            ("event_type", "VARCHAR(50)", "NOT NULL", "Sin valor", "Tipo de evento registrado."),
            ("username", "VARCHAR(100)", "-", "NULL", "Usuario asociado al evento."),
            ("ip_address", "VARCHAR(45)", "-", "NULL", "Dirección IP registrada."),
            ("success", "BOOLEAN", "-", "TRUE", "Indica si el evento fue exitoso."),
            ("message", "TEXT", "-", "NULL", "Descripción del evento."),
            ("level", "VARCHAR(10)", "-", "INFO", "Nivel de severidad."),
        ],
    ),
]


def set_tnr_run(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def paragraph_after(anchor, text=None, style=None, align=None):
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    p = Paragraph(new_p, anchor._parent)
    if style:
        p.style = style
    if text:
        run = p.add_run(text)
        set_tnr_run(run, 12)
    if align is not None:
        p.alignment = align
    return p


def add_seq_caption(anchor, label, caption, number, style="Caption"):
    p = paragraph_after(anchor, style=style, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(label + " ")
    set_tnr_run(run, 10)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), rf" SEQ {label} \* ARABIC ")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rpr.append(sz)
    t = OxmlElement("w:t")
    t.text = str(number)
    r.append(rpr)
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    run = p.add_run(": " + caption)
    set_tnr_run(run, 10)
    return p


def insert_table_after(document, anchor, rows, cols):
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    anchor._element.addnext(tbl)
    return table


def format_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_tnr_run(run, 9)
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_tnr_run(run, 9, bold=True)


def delete_between(start_p, end_p):
    cur = start_p._element.getnext()
    end_el = end_p._element
    while cur is not None and cur is not end_el:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt


def find_paragraph(doc, exact):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact:
            return paragraph
    return None


def main():
    doc = Document(DOCX)
    start = find_paragraph(doc, "5.4 Diagramas de código")
    end = find_paragraph(doc, "Capítulo 6. Desarrollo")
    if start is None or end is None:
        raise RuntimeError("No se encontró el rango 5.4 -> Capítulo 6.")

    delete_between(start, end)
    anchor = start

    for text in INTRO_54:
        anchor = paragraph_after(anchor, text, style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    anchor = paragraph_after(anchor, "5.4.1 Diagrama de módulos del prototipo", style="Título 3 TNR")
    for text in TEXT_541:
        anchor = paragraph_after(anchor, text, style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    p_img = paragraph_after(anchor, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    p_img.add_run().add_picture(str(IMG_DIR / "fig54_diagrama_modulos_codigo.png"), width=Inches(6.35))
    anchor = add_seq_caption(p_img, "Ilustración", "Diagrama de módulos de código del prototipo [autoría propia]", 52)

    anchor = paragraph_after(anchor, "5.4.2 Modelo de datos", style="Título 3 TNR")
    for text in TEXT_542:
        anchor = paragraph_after(anchor, text, style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    p_img = paragraph_after(anchor, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    p_img.add_run().add_picture(str(IMG_DIR / "fig54_modelo_entidad_relacion.png"), width=Inches(6.35))
    anchor = add_seq_caption(p_img, "Ilustración", "Modelo entidad-relación del prototipo [autoría propia]", 53)

    anchor = paragraph_after(anchor, "5.4.2.1 Diccionario de datos", style="Heading 4")
    anchor = paragraph_after(
        anchor,
        "El diccionario de datos describe los campos persistidos por el prototipo. "
        "Las tablas se derivan de schema.sql y de las migraciones usadas por la aplicación.",
        style="Normal",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    table_number = 25
    for table_name, description, rows in DICTIONARY:
        p = paragraph_after(anchor, f"Tabla “{table_name}”", style="Normal")
        set_tnr_run(p.runs[0], 12, bold=True)
        anchor = p
        anchor = paragraph_after(anchor, description, style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        table = insert_table_after(doc, anchor, len(rows) + 1, 5)
        headers = ["Campo", "Tipo de dato", "Restricciones", "Valor predeterminado", "Descripción"]
        for idx, header in enumerate(headers):
            table.cell(0, idx).text = header
        for r_idx, row_data in enumerate(rows, start=1):
            for c_idx, value in enumerate(row_data):
                table.cell(r_idx, c_idx).text = value
        format_table(table)
        anchor = Table(table._tbl, doc)
        anchor = add_seq_caption(anchor, "Tabla", f"Descripción de la tabla “{table_name}” [autoría propia]", table_number)
        table_number += 1

    anchor = paragraph_after(anchor, "Consideraciones de seguridad", style="Normal")
    set_tnr_run(anchor.runs[0], 12, bold=True)
    bullets = [
        "Las contraseñas se almacenan únicamente como hash bcrypt en password_hash.",
        "El registro se restringe a correos institucionales @ipn.mx y @alumno.ipn.mx.",
        "Los eventos de autenticación y verificación se registran en security_audit_log.",
        "Las ediciones manuales de tiempos conductuales se auditan en behavior_edits mediante snapshots before/after.",
        "El sistema opera con tres roles: admin, investigador y estudiante.",
    ]
    for item in bullets:
        p = paragraph_after(anchor, item, style="List Paragraph")
        set_tnr_run(p.runs[0], 12)
        anchor = p

    doc.save(DOCX)
    print("Sección 5.4 reemplazada correctamente.")


if __name__ == "__main__":
    main()
