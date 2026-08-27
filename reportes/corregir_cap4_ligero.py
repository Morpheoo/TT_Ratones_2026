from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")


REPLACEMENTS = {
    "Los requerimientos funcionales y no funcionales son una serie de especificaciones que determinan el funcionamiento del prototipo. La necesidad de los requerimientos funcionales es esencial para el funcionamiento del prototipo; mientras que los no funcionales a pesar de ser importantes, no son esenciales para el prototipo, pero si para su correcto funcionamiento.": (
        "Los requerimientos funcionales y no funcionales son especificaciones que determinan el funcionamiento del prototipo. Los requerimientos funcionales describen las acciones que el sistema debe ejecutar, mientras que los no funcionales establecen condiciones de calidad, seguridad, rendimiento, portabilidad y usabilidad necesarias para su operación correcta."
    ),
    "4. Unidad de Procesamiento Gráfico (GPU): GPU de arquitectura avanzada NVIDIA GeForce RTX 5070 Ti Laptop con 12 GB de VRAM dedicada, compatible con CUDA cores, Tensor cores y Compute Capability 12.0 para el entrenamiento y la inferencia rápida a más de 30 FPS del modelo YOLO11s-pose.": (
        "4. Unidad de Procesamiento Gráfico (GPU): GPU NVIDIA dedicada con soporte CUDA y al menos 8 GB de VRAM para entrenamiento e inferencia acelerada del modelo YOLO11s-pose. En el entorno de desarrollo se utilizó una NVIDIA GeForce RTX 5070 Ti Laptop con 12 GB de VRAM."
    ),
    "2. Frameworks de Aprendizaje Profundo y Visión: Librería PyTorch 2.0+ (optimizada con CUDA 12.8) y framework Ultralytics YOLO11 para la inicialización, entrenamiento y ejecución en tiempo real de la red YOLO11s-pose de keypoints.": (
        "2. Frameworks de Aprendizaje Profundo y Visión: PyTorch y Ultralytics YOLO11 para la inicialización, entrenamiento e inferencia del modelo YOLO11s-pose de keypoints, con aceleración CUDA cuando el equipo cuenta con GPU NVIDIA compatible."
    ),
    "6. Clasificadores Conductuales (LSTM y Random Forest): Modelos Random Forest (Scikit-learn) y redes neuronales recurrentes LSTM (Keras/TensorFlow) para discriminar y clasificar conductas de postura y acicalamiento dependientes del tiempo.": (
        "6. Clasificadores Conductuales (SimBA Random Forest y apoyo temporal): Modelos Random Forest mediante SimBA/scikit-learn para la clasificación productiva de Grooming y Thigmotaxis; LSTM en Keras/TensorFlow se conserva como apoyo experimental de rescate temporal, no como clasificador principal."
    ),
    "7. Interfaz Gráfica y Persistencia (Streamlit y PostgreSQL): Streamlit 1.25+ para la interfaz web interactiva y amigable del usuario, y PostgreSQL 14+ local para la persistencia segura de metadatos, ROIs y bitácoras de auditoría científica.": (
        "7. Interfaz Gráfica y Persistencia (Streamlit y PostgreSQL): Streamlit para la interfaz web interactiva del usuario, y PostgreSQL 15 local para la persistencia segura de usuarios, tratamientos, experimentos, ROIs, resultados, ediciones manuales y bitácoras de auditoría."
    ),
    "3. Ciclo III: Clasificación Conductual e Interacción (Requerimientos de Aplicación): Discriminar conductas complejas (acicalamiento, exploración) y presentarlas de forma accesible. Se definieron requerimientos para entrenar clasificadores de series temporales (LSTM) e integrar una interfaz web interactiva basada en Streamlit, eliminando la necesidad de interactuar mediante consola de comandos.": (
        "3. Ciclo III: Clasificación Conductual e Interacción (Requerimientos de Aplicación): Discriminar conductas complejas como Grooming y Thigmotaxis, además de presentar trayectoria y permanencia por zonas de forma accesible. Se consolidó el uso de SimBA Random Forest como clasificador productivo, se exploró LSTM como rescate temporal y se integró una interfaz web interactiva basada en Streamlit, eliminando la necesidad de operar mediante consola de comandos."
    ),
    "4. Ciclo IV: Seguridad y Persistencia (Requerimientos no Funcionales de Producción): Proteger los datos y posibilitar la auditoría científica. Incorporación de requerimientos de base de datos relacional (PostgreSQL) para evitar la pérdida de experimentos y el desarrollo de una bitácora de auditoría segura (behavior_edits) que documenta cada modificación manual realizada por los investigadores.": (
        "4. Ciclo IV: Seguridad y Persistencia (Requerimientos no Funcionales de Producción): Proteger los datos y posibilitar la auditoría científica. Se incorporó una base de datos relacional PostgreSQL para usuarios, tratamientos, experimentos, configuraciones ROI y resultados; además, se añadieron las tablas behavior_edits y security_audit_log para registrar ediciones manuales y eventos críticos del sistema."
    ),
    "La viabilidad algorítmica del modelo de estimación de pose se corroboró experimentalmente mediante pruebas iterativas de entrenamiento en el entorno local y en la nube. Se entrenaron y evaluaron arquitecturas YOLOv11 en su versión ligera (YOLO11s-pose), analizando las curvas de pérdida (loss curves) y la métrica de precisión media (mAP50-95) en keypoints. La convergencia del modelo se alcanzó de manera estable sin sobreajuste (overfitting). Para el análisis comportamental complementario, se modelaron y probaron clasificadores supervisados Random Forest, mediante SimBA sobre scikit-learn, y redes recurrentes LSTM implementadas en Keras/TensorFlow, comprobando la estabilidad de la clasificación ante oclusiones y oclusión parcial en el laberinto EPM.": (
        "La viabilidad algorítmica del modelo de estimación de pose se corroboró experimentalmente mediante pruebas iterativas de entrenamiento en el entorno local y en la nube. Se entrenó y evaluó YOLO11s-pose, analizando curvas de pérdida y métricas de precisión de keypoints; el modelo final alcanzó mAP50 = 0.995 con 3,953 imágenes etiquetadas. Para el análisis conductual se validó SimBA Random Forest sobre features espaciotemporales derivadas de los keypoints, y se documentaron B-SOiD y LSTM como estrategias experimentales de apoyo para casos de Grooming difíciles."
    ),
    "Esta interfaz permite cargar videos en formatos universales (MP4, AVI) sin requerir preprocesamiento manual o herramientas externas, y exportar los resultados conductuales cuantitativos directamente a archivos compatibles con Excel, asegurando la compatibilidad absoluta con el software estadístico complementario (como R, Origin o SPSS) utilizado comúnmente por los investigadores de la ENMyH-IPN.": (
        "Esta interfaz permite cargar videos en formatos comunes como MP4 y AVI sin requerir preprocesamiento manual obligatorio, y exportar los resultados conductuales cuantitativos a archivos compatibles con Excel, lo que facilita su uso posterior en herramientas estadísticas como R, Origin o SPSS."
    ),
    "4. Precisión Tolerable: Se determinó que el sistema debe alcanzar un nivel de concordancia de entre el 90% y el 95% en comparación con el análisis visual manual realizado por expertos. Este rango representa la variación estándar de tolerancia inter-observadores en el laboratorio. El prototipo cumple con este criterio mediante la calibración y el suavizado de la trayectoria con el filtro Savitzky-Golay.": (
        "4. Precisión Tolerable: Se estableció como meta operativa que las mediciones espaciales del sistema mantuvieran alta concordancia con la revisión manual, especialmente en trayectoria y permanencia por zonas. Para las conductas complejas, como Grooming, la evaluación se realizó mediante métricas de clasificación como F1-Score y validación Leave-One-Out, reconociendo que su variabilidad conductual exige más datos etiquetados y modelos temporales complementarios."
    ),
    "3. Propiedad Intelectual y Derechos de Autor: El desarrollo del prototipo (código fuente de la interfaz Streamlit, algoritmos de seguimiento y modelado temporal LSTM) es de autoría del equipo de trabajo de Trabajo Terminal y se rige por la Ley Federal del Derecho de Autor (LFDA) [109] en México. El registro de propiedad industrial e intelectual se gestionará bajo las políticas internas y el reglamento de la Dirección de Transferencia Tecnológica del IPN, asegurando la copropiedad de la institución.": (
        "3. Propiedad Intelectual y Derechos de Autor: El desarrollo del prototipo (código fuente de la interfaz Streamlit, integración YOLO-SimBA, módulos de seguimiento, persistencia y auditoría) es de autoría del equipo de Trabajo Terminal y se rige por la Ley Federal del Derecho de Autor (LFDA) [109] en México. El registro de propiedad industrial e intelectual se gestionará bajo las políticas internas y el reglamento de la Dirección de Transferencia Tecnológica del IPN, asegurando la copropiedad de la institución."
    ),
    "Procesamiento local acelerado: Gracias a la aceleración por hardware provista por los núcleos CUDA de la GPU NVIDIA GeForce RTX 5070 Ti, el prototipo realiza la estimación de postura del modelo YOLO11s-pose en tiempo real acelerado (promedio de 60 FPS). Un video típico de pruebas con una duración de 5 minutos (equivalente a 9,000 fotogramas a 30 FPS en el laberinto EPM) se procesa en un tiempo de ejecución t de solo 2.5 minutos (150 segundos, equivalente a 0.0417 horas).": (
        "Procesamiento local acelerado: Gracias a la aceleración por hardware provista por la GPU NVIDIA dedicada, el prototipo realiza la estimación de postura del modelo YOLO11s-pose en tiempo real acelerado. Un video típico de pruebas con una duración de 5 minutos (equivalente a 9,000 fotogramas a 30 FPS en el laberinto EPM) se procesa en un tiempo aproximado de 3 minutos (180 segundos, equivalente a 0.05 horas)."
    ),
    "Esto equivale a aproximadamente 2.53 gramos de CO₂ por procesamiento completo de un video de 5 minutos.": (
        "Esto equivale a aproximadamente 3.04 gramos de CO₂ por procesamiento completo de un video de 5 minutos."
    ),
    "Como se estableció previamente, el prototipo genera únicamente 0.00253 kg de CO₂ por cada procesamiento de video de 5 minutos.": (
        "Como se estableció previamente, el prototipo genera aproximadamente 0.00304 kg de CO₂ por cada procesamiento de video de 5 minutos."
    ),
    "Aun al tomar en cuenta escenarios conservadores para el transporte público y aplicar la realidad de que más de la mitad del equipo emplea este modo de transporte, el efecto medioambiental del método tradicional de análisis manual continúa siendo de 6.2 a 7.1 kg de CO₂ por jornada de experimentación (debido a los traslados obligatorios del personal); en cambio, el prototipo local optimizado solamente genera 0.00253 kg de CO₂ (2.53 g) por procesamiento completo.": (
        "Aun al tomar en cuenta escenarios conservadores para el transporte público y aplicar la realidad de que más de la mitad del equipo emplea este modo de transporte, el efecto medioambiental del método tradicional de análisis manual continúa siendo mayor por los traslados obligatorios del personal; en cambio, el prototipo local optimizado genera aproximadamente 0.00304 kg de CO₂ (3.04 g) por procesamiento completo."
    ),
    "La disminución resultante varía entre el 98.22% y el 99.96% dependiendo del escenario de transporte considerado, y alcanza un 99.96% de reducción de emisiones promedio frente al traslado obligatorio del equipo de investigadores en vehículo particular. Esto indica que el prototipo local erradica casi por completo la huella de carbono vinculada a la movilidad física, consolidándose como una opción ecológica, eficiente y sostenible alineada con las metas internacionales de desarrollo sustentable.": (
        "La disminución resultante varía de acuerdo con el escenario de transporte considerado, pero mantiene una reducción alta frente al traslado obligatorio del equipo de investigación. Esto indica que el prototipo local reduce de forma significativa la huella de carbono asociada a la movilidad física y consolida una alternativa más eficiente y sostenible para el análisis conductual."
    ),
}


def set_tnr(run, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def replace_paragraph(paragraph, text):
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    run = paragraph.add_run(text)
    set_tnr(run)


def set_cell(cell, text):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    set_tnr(run)


def main():
    doc = Document(DOCX)
    replaced = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in REPLACEMENTS:
            replace_paragraph(p, REPLACEMENTS[text])
            replaced += 1

    # Tabla 3: Requerimientos funcionales
    rf = doc.tables[6]
    set_cell(rf.cell(7, 1), "Clasificar las conductas Grooming y Thigmotaxis mediante modelos entrenados con features de postura, y calcular métricas espaciales de trayectoria por zonas.")
    set_cell(rf.cell(7, 2), "CU7 Agrupar comportamientos")

    # Tabla 4: Requerimientos no funcionales
    rnf = doc.tables[7]
    set_cell(rnf.cell(5, 1), "El sistema debe ser ejecutable en sistemas operativos Windows 10/11 mediante un entorno Python preconfigurado y dependencias documentadas.")
    set_cell(rnf.cell(5, 3), "Portabilidad mediante entorno virtual, archivo de dependencias y guía de instalación local.")

    # Tabla 12 del capítulo: sistemas operativos
    os_table = doc.tables[15]
    set_cell(os_table.cell(1, 4), "Seleccionado por su amplia difusión en el laboratorio y compatibilidad con GPU NVIDIA, Python, Streamlit y PostgreSQL.")

    # Tabla de estrategias de mitigación
    risk = doc.tables[22]
    set_cell(risk.cell(2, 4), "Restauración del entorno Python mediante requirements.txt, guía de instalación y respaldo de configuración.")
    set_cell(risk.cell(7, 2), "Curación del dataset, validación Leave-One-Out, uso de SimBA Random Forest y estrategia Conditional para Grooming.")
    set_cell(risk.cell(7, 3), "Revisión etológica de videos críticos, ajuste de umbrales y documentación de casos ambiguos.")
    set_cell(risk.cell(7, 4), "Auditoría y corrección manual trazable mediante behavior_edits cuando el modelo no capture una conducta compleja.")

    doc.save(DOCX)
    print(f"Parrafos reemplazados: {replaced}")
    print("Tablas actualizadas: RF7, RNF5, SO Windows, R02 y R07.")


if __name__ == "__main__":
    main()
