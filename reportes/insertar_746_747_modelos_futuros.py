from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")


BLOCKS = [
    (
        "7.4.6 Estimación del tamaño de dataset requerido",
        [
            (
                "El principal factor limitante para mejorar la precisión de Grooming no es el "
                "rastreador de pose, sino la variabilidad conductual disponible para entrenar el "
                "clasificador. A diferencia de Thigmotaxis, que depende en mayor medida de la "
                "posición del centroide respecto a la geometría del laberinto, Grooming depende de "
                "micro-movimientos repetitivos de cabeza, patas y cuerpo. Por ello, requiere mayor "
                "diversidad de animales, orientaciones, duraciones y estilos de acicalamiento."
            ),
            (
                "La estimación del número de videos necesarios puede aproximarse a partir del "
                "intervalo de confianza del promedio de F1-Score observado en la validación "
                "Leave-One-Out. Para Grooming, la estrategia Conditional obtuvo F1-Score justo de "
                "0.523 con desviación estándar aproximada de 0.315 en n efectivo = 24. Si se busca "
                "estimar el F1 promedio con un margen de error E cercano a +/-0.06 y un nivel de "
                "confianza del 95 %, se puede usar la expresión n = (1.96 sigma / E)^2. Al sustituir "
                "sigma = 0.315 y E = 0.06, se obtiene n ≈ 106 videos."
            ),
            (
                "Con base en esta aproximación, una meta razonable para futuras iteraciones es "
                "ampliar el conjunto a 100-120 videos etiquetados. Este rango no debe interpretarse "
                "como un umbral absoluto definido por un algoritmo externo, sino como una estimación "
                "estadística derivada de la varianza observada en este proyecto. Un conjunto de ese "
                "tamaño permitiría reducir la incertidumbre del promedio, mejorar la cobertura de "
                "casos poco frecuentes y entrenar modelos temporales o no supervisados con menor "
                "riesgo de sobreajuste."
            ),
            (
                "Esta estimación también explica por qué Thigmotaxis puede estabilizarse antes que "
                "Grooming. La primera conducta se relaciona con una condición espacial relativamente "
                "estable, mientras que Grooming requiere reconocer secuencias cinemáticas finas. "
                "B-SOiD reporta que los patrones de acicalamiento pueden separarse en motivos como "
                "paw/face groom, head groom y body lick, lo que refuerza que una sola etiqueta de "
                "Grooming agrupa subconductas con dinámica interna distinta [116]."
            ),
        ],
    ),
    (
        "7.4.7 Modelos candidatos para mejorar la precisión en trabajo futuro",
        [
            (
                "A partir de los resultados del prototipo y de la literatura reciente en análisis "
                "automatizado de conducta, el trabajo futuro debe priorizar modelos que aprovechen "
                "la estructura temporal de los keypoints, reduzcan el costo de etiquetado y permitan "
                "distinguir subpatrones de Grooming. Las rutas más relevantes son las siguientes:"
            ),
            (
                "Keypoint-MoSeq. Es el candidato más sólido para explorar Grooming con mayor detalle, "
                "porque modela sílabas conductuales a partir de trayectorias de keypoints y separa "
                "ruido de cambios reales en la dinámica corporal. En reportes recientes, Keypoint-MoSeq "
                "ha mostrado mejor alineación con transiciones conductuales que alternativas como "
                "B-SOiD, VAME y MotionMapper, además de recuperar estados relacionados con grooming "
                "en datos de campo abierto [117]. Para este proyecto sería útil como herramienta de "
                "descubrimiento de subconductas antes de entrenar un clasificador supervisado final."
            ),
            (
                "B-SOiD refinado. Aunque en la validación actual no superó a Conditional, sigue siendo "
                "útil como apoyo no supervisado para descubrir motivos cinemáticos de Grooming. Con "
                "más videos, podría reentrenarse para mapear motivos específicos a Grooming real y "
                "usar sus salidas como variables adicionales del clasificador principal. Esta ruta "
                "es defendible porque B-SOiD fue diseñado para identificar patrones espaciotemporales "
                "a partir de poses y luego entrenar un Random Forest sobre esos patrones [116]."
            ),
            (
                "Aprendizaje activo. En lugar de etiquetar videos al azar, una siguiente versión puede "
                "seleccionar fragmentos donde el modelo muestre mayor incertidumbre o desacuerdo entre "
                "clasificadores. Esto concentraría el esfuerzo humano en los casos que más información "
                "aportan: transiciones ambiguas, grooming parcial, grooming breve y posturas poco "
                "representadas. Esta estrategia es especialmente conveniente cuando el etiquetado frame "
                "a frame es costoso."
            ),
            (
                "Modelos temporales supervisados. Con un dataset ampliado, pueden evaluarse redes "
                "temporales como GRU, LSTM o Temporal Convolutional Networks sobre ventanas de keypoints "
                "y features SimBA. Estas arquitecturas podrían mejorar Grooming porque clasifican la "
                "evolución de una secuencia y no únicamente el estado local de un frame. La capa LSTM "
                "ya probada en este trabajo debe conservarse como antecedente experimental, pero una "
                "validación ciega completa requeriría mayor número de videos y separación estricta por "
                "animal."
            ),
            (
                "Modelos de video supervisados. Herramientas como DeepEthogram y MARS muestran que la "
                "clasificación directa desde video puede ser útil para conductas visualmente complejas "
                "[119], [120]. En este proyecto podrían explorarse como comparación futura, aunque su "
                "costo computacional y su menor explicabilidad los vuelven menos convenientes como "
                "primer reemplazo del flujo YOLO Pose + SimBA."
            ),
            (
                "Postprocesamiento probabilístico. Un modelo de Markov oculto o un campo aleatorio "
                "condicional podría mejorar la coherencia temporal de las predicciones sin reemplazar "
                "al clasificador base. Este enfoque reduciría segmentos aislados de uno o pocos frames "
                "y ayudaría a respetar duraciones plausibles de cada conducta."
            ),
            (
                "Con base en lo anterior, la ruta técnica recomendada es conservar SimBA Random Forest "
                "como línea base explicable, ampliar el dataset a 100-120 videos, usar Keypoint-MoSeq "
                "o B-SOiD para descubrir subconductas de Grooming y, posteriormente, entrenar un modelo "
                "temporal supervisado sobre ventanas de movimiento. Esta estrategia mantiene la "
                "interpretabilidad del sistema actual, pero abre una ruta concreta para mejorar la "
                "precisión en especímenes no vistos."
            ),
        ],
    ),
]

REFERENCES = [
    (
        "[116] A. I. Hsu y E. A. Yttri, “B-SOiD, an open-source unsupervised algorithm "
        "for identification and fast prediction of behaviors,” Nature Communications, vol. 12, "
        "art. 5188, 2021. doi: 10.1038/s41467-021-25420-x."
    ),
    (
        "[117] C. Weinreb, J. Pearl, S. Lin et al., “Keypoint-MoSeq: parsing behavior by linking "
        "point tracking to pose dynamics,” Nature Methods, vol. 21, pp. 1329-1339, 2024. "
        "doi: 10.1038/s41592-024-02318-2."
    ),
    (
        "[118] N. L. Goodwin, J. J. Choong, S. Hwang et al., “Simple Behavioral Analysis "
        "(SimBA) as a platform for explainable machine learning in behavioral neuroscience,” "
        "Nature Neuroscience, vol. 27, pp. 1411-1424, 2024. doi: 10.1038/s41593-024-01649-9."
    ),
    (
        "[119] C. Segalin et al., “The Mouse Action Recognition System (MARS),” eLife, vol. 10, "
        "e63720, 2021. doi: 10.7554/eLife.63720."
    ),
    (
        "[120] J. P. Bohnslav et al., “DeepEthogram, a machine learning pipeline for supervised "
        "behavior classification from raw pixels,” eLife, vol. 10, e63377, 2021. "
        "doi: 10.7554/eLife.63377."
    ),
    (
        "[121] K. Luxem et al., “Identifying behavioral structure from deep variational embeddings,” "
        "Communications Biology, vol. 5, art. 1267, 2022. doi: 10.1038/s42003-022-04080-7."
    ),
]


def set_tnr(run, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_after(anchor, text, style):
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    p = Paragraph(new_p, anchor._parent)
    p.style = style
    run = p.add_run(text)
    set_tnr(run)
    return p


def find_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def paragraph_exists(doc, text):
    return any(p.text.strip() == text for p in doc.paragraphs)


def main():
    doc = Document(DOCX)

    conclusions = find_paragraph(doc, "Conclusiones")
    if conclusions is None:
        raise RuntimeError("No se encontró el encabezado 'Conclusiones'.")
    if paragraph_exists(doc, "7.4.6 Estimación del tamaño de dataset requerido"):
        raise RuntimeError("La sección 7.4.6 ya existe; no se insertó de nuevo.")

    previous = conclusions._element.getprevious()
    if previous is None:
        raise RuntimeError("No se encontró el párrafo anterior a Conclusiones.")
    anchor = Paragraph(previous, conclusions._parent)

    for heading, paragraphs in BLOCKS:
        anchor = add_after(anchor, heading, "Título 3 TNR")
        for text in paragraphs:
            anchor = add_after(anchor, text, "Texto Normal TNR")

    last_ref = find_paragraph(doc, "[115] Banco de México, “Tipo de cambio FIX, 5 de junio de 2026: 17.4755 MXN/USD,” Portal del Mercado Cambiario, 2026. [En línea]. Disponible en: https://www.banxico.org.mx/portal-mercado-cambiario/ (consultado el 7 de junio de 2026).")
    if last_ref is None:
        raise RuntimeError("No se encontró la referencia [115] para anexar nuevas referencias.")

    anchor = last_ref
    for ref in REFERENCES:
        if not paragraph_exists(doc, ref):
            anchor = add_after(anchor, ref, "Texto Normal TNR")

    doc.save(DOCX)
    print("Subsecciones 7.4.6 y 7.4.7 insertadas; referencias [116]-[121] agregadas.")


if __name__ == "__main__":
    main()
