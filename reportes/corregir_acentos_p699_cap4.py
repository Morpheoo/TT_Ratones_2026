from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")


TEXT = (
    "El prototipo fue dise\u00f1ado bajo una arquitectura de software modular e independiente, "
    "facilitando su integraci\u00f3n con los sistemas operativos de escritorio est\u00e1ndar en los "
    "laboratorios (Windows 10/11). La persistencia de datos se resolvi\u00f3 mediante una base de datos "
    "relacional PostgreSQL local, la cual se comunica con la interfaz interactiva desarrollada en "
    "Streamlit. Esta interfaz permite cargar videos en formatos comunes como MP4 y AVI sin requerir "
    "preprocesamiento manual obligatorio, y exportar los resultados conductuales cuantitativos a "
    "archivos compatibles con Excel, lo que facilita su uso posterior en herramientas estad\u00edsticas "
    "como R, Origin o SPSS."
)


def set_tnr(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def main():
    doc = Document(DOCX)
    count = 0
    for p in doc.paragraphs:
        if "dise?ado" in p.text and "compatibles con Excel" in p.text:
            style = p.style
            p.clear()
            p.style = style
            run = p.add_run(TEXT)
            set_tnr(run)
            count += 1
    doc.save(DOCX)
    print(f"Parrafos corregidos: {count}")


if __name__ == "__main__":
    main()
