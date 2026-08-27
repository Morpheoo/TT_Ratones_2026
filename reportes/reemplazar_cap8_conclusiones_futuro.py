from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")


BLOCKS = [
    (
        "Capítulo 8. Conclusiones",
        "Título 1 TNR",
        [],
    ),
    (
        "8.1 Conclusiones generales",
        "Título 2 TNR",
        [
            (
                "El desarrollo del prototipo permitió comprobar la viabilidad técnica de un sistema "
                "local de visión por computadora para apoyar el análisis de conducta en roedores dentro "
                "del modelo Elevated Plus Maze. La solución integra estimación de pose markerless, "
                "extracción de características espaciotemporales, clasificación conductual, edición "
                "manual trazable, persistencia relacional y exportación de resultados en un flujo "
                "operativo accesible para usuarios de laboratorio."
            ),
            (
                "Uno de los resultados más sólidos del proyecto fue la sustitución de enfoques de pose "
                "más costosos en tiempo de ejecución por YOLO11 Pose mediante Ultralytics. El modelo "
                "entrenado con 3,953 imágenes etiquetadas alcanzó mAP50 = 0.995 y permitió procesar "
                "videos de 5 minutos en aproximadamente 3 minutos con GPU dedicada. Esta mejora hizo "
                "posible que el análisis automatizado se ejecutara en una estación local, sin depender "
                "de servicios externos ni de flujos manuales prolongados."
            ),
            (
                "En la clasificación conductual, el proyecto mostró que Thigmotaxis y Grooming tienen "
                "niveles de dificultad distintos. Thigmotaxis se benefició de su naturaleza espacial y "
                "alcanzó mejor estabilidad con SimBA Random Forest. Grooming presentó mayor variabilidad "
                "por tratarse de una conducta cinemática fina, con movimientos breves, repetitivos y "
                "dependientes de postura. Por esta razón, el sistema final no se presenta como un único "
                "clasificador perfecto, sino como una arquitectura comparada y justificable: SimBA "
                "Random Forest como base productiva, estrategia Conditional para mejorar Grooming y "
                "LSTM como apoyo experimental de rescate temporal."
            ),
            (
                "La validación Leave-One-Out con 26 videos permitió documentar el desempeño real del "
                "sistema sin ocultar sus limitaciones. Grooming obtuvo su mejor resultado con la "
                "estrategia Conditional, con F1-Score justo de 0.523, mientras que Thigmotaxis alcanzó "
                "F1-Score justo de 0.636 con SimBA Random Forest. Estos valores no deben interpretarse "
                "como cierre definitivo del problema, sino como evidencia de un prototipo funcional, "
                "medible y mejorable con mayor diversidad de datos etiquetados."
            ),
            (
                "Además del componente de inteligencia artificial, el proyecto consolidó una base "
                "operativa de software: manejo de usuarios por roles, registro de tratamientos y "
                "experimentos, configuración de regiones de interés, almacenamiento en PostgreSQL, "
                "bitácora de auditoría de seguridad y registro de ediciones manuales mediante "
                "behavior_edits. Esta trazabilidad es relevante porque permite que los resultados "
                "automatizados sean revisables por investigadores y no funcionen como una caja negra."
            ),
        ],
    ),
    (
        "8.2 Cumplimiento de objetivos",
        "Título 2 TNR",
        [
            (
                "El objetivo general se cumplió al desarrollar un prototipo capaz de procesar videos "
                "del modelo EPM, estimar keypoints del espécimen, calcular trayectoria, clasificar "
                "conductas de interés y presentar resultados consultables por el usuario. La interfaz "
                "en Streamlit redujo la dependencia de consola y permitió organizar el flujo de trabajo "
                "en etapas comprensibles: registro, configuración de zonas, análisis, consulta, edición "
                "y exportación."
            ),
            (
                "También se cumplieron objetivos técnicos específicos. Se entrenó y validó el modelo "
                "YOLO11 Pose; se construyó el puente de features hacia SimBA; se entrenaron modelos "
                "Random Forest para Grooming y Thigmotaxis; se evaluaron estrategias alternativas como "
                "B-SOiD, Ensemble OR, Dynamic y LSTM rescue; y se documentó una arquitectura final basada "
                "en evidencia experimental. La base de datos quedó estructurada para conservar usuarios, "
                "tratamientos, experimentos, configuraciones ROI, resultados, ediciones y auditoría."
            ),
            (
                "Desde el punto de vista metodológico, el uso del modelo en espiral permitió iterar sobre "
                "riesgos reales del proyecto: rendimiento de inferencia, ruido en keypoints, variabilidad "
                "de conducta, portabilidad del entorno y trazabilidad de modificaciones. Esta forma de "
                "trabajo fue adecuada para un prototipo científico, donde los requerimientos evolucionaron "
                "conforme se obtuvieron resultados experimentales."
            ),
        ],
    ),
    (
        "8.3 Limitaciones identificadas",
        "Título 2 TNR",
        [
            (
                "La principal limitación del sistema se encuentra en el tamaño y diversidad del dataset "
                "conductual etiquetado. Aunque se contó con 26 videos reales y 243,253 frames, Grooming "
                "representó solo 20,757 frames positivos, equivalentes al 8.5 % del total. Esta baja "
                "prevalencia incrementa la dificultad de aprendizaje y explica la variabilidad observada "
                "entre especímenes durante la validación Leave-One-Out."
            ),
            (
                "Otra limitación es la diferencia entre conductas espaciales y cinemáticas. Thigmotaxis "
                "puede inferirse con mayor estabilidad a partir de posición y regiones de interés; en "
                "cambio, Grooming depende de micro-movimientos y de la continuidad temporal de secuencias "
                "breves. Por ello, un clasificador tabular como Random Forest puede requerir apoyo de "
                "modelos temporales o métodos de descubrimiento de motivos para mejorar su sensibilidad."
            ),
            (
                "La portabilidad del prototipo también requiere mantenimiento. Aunque el sistema opera "
                "en Windows y se realizaron pruebas de instalación en Linux, las dependencias de visión "
                "por computadora, SimBA, PyTorch, TensorFlow/Keras y Ultralytics pueden generar conflictos "
                "entre versiones. Por esta razón, la documentación de instalación y los entornos virtuales "
                "deben mantenerse actualizados."
            ),
            (
                "Finalmente, el prototipo no sustituye la validación experta. Su propósito es reducir el "
                "tiempo de análisis, estandarizar mediciones y apoyar la revisión conductual, pero las "
                "conductas ambiguas deben conservar mecanismos de edición manual, justificación y auditoría "
                "para mantener validez científica."
            ),
        ],
    ),
    (
        "8.4 Trabajo a futuro",
        "Título 2 TNR",
        [
            (
                "La mejora prioritaria consiste en ampliar el conjunto de videos etiquetados. Con base "
                "en la varianza observada durante la validación Leave-One-Out, se estimó que para "
                "Grooming una meta razonable se ubica entre 100 y 120 videos etiquetados. Este rango "
                "permitiría reducir la incertidumbre del F1 promedio, cubrir más estilos de acicalamiento "
                "y entrenar modelos con mayor capacidad de generalización."
            ),
            (
                "Una segunda línea de trabajo consiste en explorar modelos especializados para conducta "
                "temporal. Keypoint-MoSeq representa una alternativa fuerte para descubrir subconductas "
                "o sílabas conductuales a partir de keypoints; B-SOiD puede refinarse con más videos para "
                "identificar motivos de Grooming; y modelos supervisados como GRU, LSTM o Temporal "
                "Convolutional Networks pueden evaluarse cuando exista un dataset más amplio y balanceado."
            ),
            (
                "También se recomienda incorporar aprendizaje activo para reducir el costo de etiquetado. "
                "En lugar de etiquetar videos al azar, el sistema podría seleccionar fragmentos donde el "
                "clasificador tenga mayor incertidumbre o donde exista desacuerdo entre modelos. Esto "
                "concentraría el esfuerzo humano en los segmentos que más información aportan al entrenamiento."
            ),
            (
                "En términos de software, el trabajo futuro debe fortalecer la portabilidad del sistema "
                "mediante entornos reproducibles, pruebas de instalación automatizadas y documentación "
                "diferenciada para Windows y Linux. También conviene mantener la compatibilidad de modelos "
                "pesados fuera del repositorio mediante rutas configurables, verificación de integridad y "
                "procedimientos claros de descarga o transferencia."
            ),
            (
                "Finalmente, se propone ampliar la validación con nuevos especímenes, sesiones y condiciones "
                "experimentales, así como comparar los resultados del prototipo contra herramientas comerciales "
                "o protocolos manuales estandarizados. Esta etapa permitiría transformar el sistema de un "
                "prototipo funcional de Trabajo Terminal a una herramienta científica más robusta para uso "
                "institucional."
            ),
        ],
    ),
]


def set_tnr(run, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def find(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def add_after(anchor, text, style):
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    p = Paragraph(new_p, anchor._parent)
    p.style = style
    run = p.add_run(text)
    set_tnr(run)
    return p


def delete_between(start, end):
    cur = start._element.getnext()
    end_el = end._element
    while cur is not None and cur is not end_el:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt


def replace_paragraph(paragraph, text, style):
    paragraph.clear()
    paragraph.style = style
    run = paragraph.add_run(text)
    set_tnr(run)


def main():
    doc = Document(DOCX)
    start = find(doc, "Conclusiones")
    end = find(doc, "Referencias")
    if start is None or end is None:
        raise RuntimeError("No se encontraron los encabezados Conclusiones/Referencias.")

    delete_between(start, end)
    first_heading, first_style, _ = BLOCKS[0]
    replace_paragraph(start, first_heading, first_style)

    anchor = start
    for heading, style, paragraphs in BLOCKS[1:]:
        anchor = add_after(anchor, heading, style)
        for text in paragraphs:
            anchor = add_after(anchor, text, "Texto Normal TNR")

    doc.save(DOCX)
    print("Capítulo 8 reemplazado.")


if __name__ == "__main__":
    main()
