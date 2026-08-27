from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")


def set_tnr(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def replace(p, text):
    style = p.style
    p.clear()
    p.style = style
    run = p.add_run(text)
    set_tnr(run)


def main():
    doc = Document(DOCX)
    replaced = 0
    for p in doc.paragraphs:
        t = p.text.strip()

        if t.startswith("Como respuesta directa a la alta demanda computacional detectada"):
            replace(p,
                "Como respuesta a la demanda computacional del análisis de video, el proyecto migró "
                "hacia inferencia acelerada por GPU para ejecutar YOLO11 Pose de forma local. Esta "
                "decisión redujo el cuello de botella observado en etapas iniciales y permitió procesar "
                "videos de 5 minutos en aproximadamente 3 minutos con GPU dedicada, manteniendo la "
                "integridad espacial necesaria para el análisis de trayectoria y keypoints.")
            replaced += 1

        elif t.startswith("Este avance técnico se complementó con la implementación de métodos heurísticos"):
            replace(p,
                "La arquitectura de ejecución se organizó mediante entornos Python separados para "
                "mantener compatibilidad entre el pipeline de visión por computadora, SimBA, "
                "TensorFlow/Keras y las dependencias de video. Esta separación redujo conflictos de "
                "versiones y permitió aislar los módulos de entrenamiento, inferencia y análisis "
                "conductual.")
            replaced += 1

        elif t.startswith("La robustez de la infraestructura seleccionada permitió además sostener ciclos"):
            replace(p,
                "La infraestructura local fue suficiente para entrenar y ejecutar el modelo YOLO11 "
                "Pose usado en el prototipo. El modelo final se entrenó con 3,953 imágenes etiquetadas "
                "y se integró al flujo YOLO-SimBA para generar keypoints, trayectoria y features "
                "espaciotemporales utilizadas por los clasificadores conductuales.")
            replaced += 1

        elif t.startswith("Para mitigar el riesgo de inconsistencia, se estableció que el sistema debe operar exclusivamente"):
            replace(p,
                "Para mitigar el riesgo de inconsistencia, se estableció que el sistema debe trabajar "
                "con video en resolución nativa 720p siempre que sea posible. La reducción de resolución "
                "incrementa el jitter de keypoints y afecta directamente las features utilizadas por "
                "SimBA. En la versión final, DeepLabCut se conserva como antecedente metodológico, pero "
                "la estimación de pose productiva se realiza con YOLO11 Pose.")
            replaced += 1

        elif t.startswith("Se determinó que el sacrificio de resolución invalida la precisión del modelo etológico"):
            replace(p,
                "Se determinó que la optimización de velocidad debe realizarse mediante hardware, "
                "postprocesamiento y modelos ligeros, no mediante degradación del material visual. Esta "
                "decisión favorece trayectorias más estables y reduce falsos positivos derivados de "
                "ruido espacial.")
            replaced += 1

        elif t.startswith("La consolidación del motor de detección y seguimiento transitó"):
            replace(p,
                "La consolidación del motor de detección transitó por varias iteraciones de entrenamiento "
                "hasta llegar al modelo YOLO11 Pose final. Las versiones tempranas tuvieron problemas de "
                "generalización ante cambios de iluminación, contraste y presencia del operador en escena. "
                "La curación del dataset y el aumento progresivo de imágenes permitieron reducir estos "
                "fallos.")
            replaced += 1

        elif t.startswith("El análisis detallado permitió identificar que el modelo confundía"):
            replace(p,
                "El análisis detallado permitió identificar errores de confusión entre el espécimen y "
                "elementos externos de la escena, especialmente cuando aparecía personal de laboratorio. "
                "La versión final se entrenó con 3,953 imágenes etiquetadas para mejorar la robustez ante "
                "variaciones de iluminación, orientación corporal y contraste del laberinto.")
            replaced += 1

        elif t.startswith("Las métricas de desempeño obtenidas tras este entrenamiento consolidado"):
            replace(p,
                "El modelo final de pose alcanzó mAP50 = 0.995 en el conjunto de validación interno. "
                "Este valor corresponde al desempeño del detector de keypoints y no debe confundirse con "
                "las métricas de clasificación conductual. Para Grooming y Thigmotaxis, el desempeño se "
                "reporta mediante F1-Score en la validación Leave-One-Out descrita en el capítulo 7.")
            replaced += 1

        elif t.startswith("El éxito del modelo V3 también radica"):
            replace(p,
                "La versión operativa del prototipo quedó basada en YOLO11 Pose v4 para estimación de "
                "keypoints. Esta decisión permitió estabilizar el seguimiento espacial y alimentar el "
                "puente de features hacia SimBA, que posteriormente clasifica Grooming y Thigmotaxis.")
            replaced += 1

        elif t.startswith("Por otro lado, la herramienta DeepLabCut"):
            replace(p,
                "Por otro lado, SimBA ejecuta la clasificación conductual a partir de features derivadas "
                "de los keypoints generados por YOLO11 Pose. DeepLabCut ya no forma parte del pipeline "
                "productivo; se conserva únicamente como antecedente comparativo de estimación de pose "
                "animal. El sistema utiliza umbrales calibrados, suavizado temporal y estrategia Conditional "
                "para reducir eventos fragmentados en Grooming.")
            replaced += 1

        elif t.startswith("La capacidad de generalización del sistema se verificó mediante la ejecución de una"):
            replace(p,
                "La capacidad de generalización del sistema se evaluó mediante validación Leave-One-Out "
                "sobre 26 videos reales etiquetados. Este enfoque permitió medir el desempeño en videos no "
                "vistos durante entrenamiento y documentar diferencias entre conductas espaciales, como "
                "Thigmotaxis, y conductas cinemáticas finas, como Grooming.")
            replaced += 1

        elif t.startswith("Este error de identificación, que dispara falsos positivos"):
            replace(p,
                "Los errores observados en videos críticos confirmaron la importancia de conservar "
                "mecanismos de revisión manual y auditoría. En lugar de presentar una concordancia absoluta, "
                "el sistema reporta métricas por conducta y permite registrar correcciones trazables cuando "
                "el modelo no captura correctamente eventos ambiguos.")
            replaced += 1

        elif t.startswith("Para mitigar los riesgos de saturación por datos de baja calidad"):
            replace(p,
                "Para mitigar riesgos de sobreajuste y baja representatividad, el proyecto consolidó un "
                "dataset de pose de 3,953 imágenes y un conjunto conductual de 26 videos reales con "
                "243,253 frames. Esta base permitió entrenar el detector de pose y evaluar los "
                "clasificadores conductuales, aunque Grooming requiere ampliación futura del dataset.")
            replaced += 1

        elif t.startswith("Durante el transcurso de la semana 42"):
            replace(p,
                "Durante las iteraciones finales se consolidó el dataset de estimación de pose hasta "
                "alcanzar 3,953 imágenes etiquetadas. Este conjunto reemplaza los cortes parciales "
                "reportados en ciclos previos y constituye la base canónica del modelo YOLO11 Pose "
                "utilizado en el prototipo.")
            replaced += 1

        elif t.startswith("El objetivo técnico de integrar estas 4,170 nuevas muestras"):
            replace(p,
                "El objetivo técnico de ampliar y depurar el dataset fue mejorar la robustez del detector "
                "ante distintas orientaciones, condiciones de iluminación y posturas del espécimen. Para "
                "la clasificación de microconductas, el proyecto utiliza features de pose en SimBA y deja "
                "como trabajo futuro la ampliación del conjunto conductual a 100-120 videos etiquetados.")
            replaced += 1

        elif t.startswith("En concordancia con los hallazgos críticos de las semanas 35 y 36"):
            replace(p,
                "En concordancia con los hallazgos de jitter, se mantuvo el criterio de trabajar con "
                "resolución nativa 720p y evitar degradaciones que afecten los keypoints. Esta decisión "
                "reduce errores de trayectoria y favorece features más estables para SimBA.")
            replaced += 1

        elif t.startswith("Paralelamente a las tareas de etiquetado masivo"):
            replace(p,
                "Paralelamente a la curación del dataset, se realizaron pruebas de integración entre "
                "YOLO11 Pose, el puente de features, SimBA, PostgreSQL y la interfaz Streamlit. Estas "
                "pruebas confirmaron la estabilidad del flujo principal y ayudaron a detectar problemas "
                "de rutas, dependencias y modelos pesados.")
            replaced += 1

        elif t.startswith("En cuanto al enriquecimiento del motor de inteligencia artificial"):
            replace(p,
                "En cuanto al enriquecimiento del motor de inteligencia artificial, el dataset final de "
                "pose quedó consolidado en 3,953 imágenes etiquetadas. Los conjuntos parciales de "
                "iteraciones anteriores se integraron y depuraron para entrenar el modelo YOLO11 Pose "
                "utilizado en el prototipo.")
            replaced += 1

        elif t.startswith("Desde una perspectiva técnica de implementación, se identificó un conflicto"):
            replace(p,
                "Desde una perspectiva técnica de implementación, se identificaron diferencias entre "
                "formatos de anotación y exportación de modelos de pose. La solución final mantuvo un "
                "esquema consistente de keypoints anatómicos y rutas de inferencia compatibles con el "
                "pipeline YOLO-SimBA.")
            replaced += 1

        elif t.startswith("Como parte de la estrategia de mejora continua, se ha mantenido"):
            replace(p,
                "Como parte de la estrategia de mejora continua, el trabajo futuro ya no se plantea como "
                "una meta de 8,000 imágenes de pose, sino como la ampliación del dataset conductual. Con "
                "base en la varianza observada en Grooming, se recomienda reunir aproximadamente 100-120 "
                "videos etiquetados para reducir la incertidumbre del F1 promedio y entrenar modelos "
                "temporales con mayor capacidad de generalización.")
            replaced += 1

        elif t.startswith("Finalmente, el éxito de estas demostraciones"):
            replace(p,
                "Finalmente, estas demostraciones permitieron validar el flujo operativo del prototipo "
                "ante necesidades reales del laboratorio: selección de zonas, carga de video, inferencia "
                "YOLO11 Pose, clasificación con SimBA, consulta de resultados y exportación estructurada. "
                "El sistema quedó preparado para continuar la validación con nuevos videos y para "
                "documentar cuantitativamente sus resultados mediante métricas de clasificación.")
            replaced += 1

        elif t.startswith("El desarrollo tecnológico del sistema ha alcanzado un hito crítico"):
            replace(p,
                "El desarrollo tecnológico del sistema alcanzó un hito crítico con la consolidación del "
                "pipeline YOLO11 Pose + SimBA. Esta integración evita depender de DeepLabCut como backend "
                "productivo y conserva la resolución nativa de 720p para reducir jitter en los keypoints "
                "que alimentan las features conductuales.")
            replaced += 1

        elif t.startswith("Para fortalecer la robustez del rastreo anatómico, el Commit 917b66d"):
            replace(p,
                "Para fortalecer la robustez del rastreo anatómico, el pipeline utiliza las cajas y "
                "keypoints generados por YOLO11 Pose, junto con filtros de suavizado y validación de "
                "trayectoria. Esta estrategia reduce falsos positivos y mantiene la coherencia espacial "
                "del espécimen en zonas de bajo contraste.")
            replaced += 1

        elif t.startswith("Finalmente, la efectividad del sistema para la toma de decisiones farmacológicas"):
            replace(p,
                "Finalmente, la utilidad del sistema para apoyar decisiones experimentales se verificó "
                "mediante reportes estructurados de trayectoria, permanencia por zonas y tiempos "
                "conductuales. La interpretación farmacológica final permanece a cargo de los especialistas, "
                "mientras que el prototipo aporta mediciones estandarizadas y trazables.")
            replaced += 1

        elif t.startswith("En el ámbito de la integridad de los metadatos etológicos"):
            replace(p,
                "En el ámbito de la integridad de los metadatos etológicos, se ajustó el flujo de "
                "sincronización entre frames de video, keypoints y features tabulares. Esta alineación "
                "temporal es necesaria para que SimBA procese secuencias completas y para evitar sesgos "
                "en la interpretación de Thigmotaxis y Grooming.")
            replaced += 1

        elif t.startswith("Como parte del esfuerzo continuo de enriquecimiento de los modelos de inteligencia artificial"):
            replace(p,
                "Como parte del esfuerzo continuo de enriquecimiento de los modelos de inteligencia "
                "artificial, se consolidó la anotación de keypoints del espécimen en RoboFlow hasta "
                "integrar el dataset final de 3,953 imágenes. Estas muestras alimentaron el entrenamiento "
                "del modelo YOLO11 Pose utilizado como backend de estimación de postura.")
            replaced += 1

        elif t.startswith("La consolidación de estos avances permite que el prototipo no solo sea"):
            replace(p,
                "La consolidación de estos avances permitió que el prototipo alcanzara un estado funcional "
                "para validación académica. La eliminación de errores de ruta, la integración con base de "
                "datos y la estandarización de exportaciones reducen riesgos operativos durante el análisis "
                "de nuevos videos experimentales.")
            replaced += 1

        elif t.startswith("Prototipo Il. En el segundo prototipo"):
            replace(p,
                "Prototipo II. En el segundo prototipo se logró un avance relevante al incrementar el "
                "conjunto de imágenes de entrenamiento respecto a la versión inicial. Esta etapa permitió "
                "identificar problemas de confusión con el operador y motivó la curación posterior del "
                "dataset hasta consolidar el modelo final YOLO11 Pose con 3,953 imágenes etiquetadas.")
            replaced += 1

    doc.save(DOCX)
    print(f"Parrafos saneados: {replaced}")


if __name__ == "__main__":
    main()
