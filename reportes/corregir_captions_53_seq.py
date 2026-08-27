from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")

CAPTIONS = [
    "Diagrama de secuencia CU1: Registrar usuario",
    "Diagrama de secuencia CU2: Iniciar sesión",
    "Diagrama de secuencia CU3: Gestionar tratamientos",
    "Diagrama de secuencia CU4: Registrar experimento",
    "Diagrama de secuencia CU5: Configurar ROIs",
    "Diagrama de secuencia CU6: Ejecutar análisis de video",
    "Diagrama de secuencia CU7: Agrupar comportamientos",
    "Diagrama de secuencia CU8: Editar tiempos conductuales",
    "Diagrama de secuencia CU9: Exportar resultados",
    "Diagrama de secuencia CU10: Consultar bitácora de auditoría",
]


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_tnr_run(paragraph, text="", size=10):
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    return run


def add_seq_field(paragraph, cached_number):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r" SEQ Ilustración \* ARABIC ")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rpr.append(sz)
    text = OxmlElement("w:t")
    text.text = str(cached_number)
    run.append(rpr)
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def main():
    doc = Document(DOCX)
    capture = False
    caption_idx = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "5.3 Diagramas de secuencia":
            capture = True
            continue
        if capture and text == "5.4 Diagramas de código":
            break
        if capture and paragraph.style.name == "Caption" and caption_idx < len(CAPTIONS):
            clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_tnr_run(paragraph, "Ilustración ")
            add_seq_field(paragraph, 42 + caption_idx)
            add_tnr_run(paragraph, f": {CAPTIONS[caption_idx]} [autoría propia]")
            caption_idx += 1

    if caption_idx != len(CAPTIONS):
        raise RuntimeError(f"Solo se actualizaron {caption_idx} captions de {len(CAPTIONS)}.")

    doc.save(DOCX)
    print("Captions de 5.3 corregidos con campos SEQ.")


if __name__ == "__main__":
    main()
