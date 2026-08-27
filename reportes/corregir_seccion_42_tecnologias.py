from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")


PARAGRAPH_REPLACEMENTS = {
    "Las herramientas por utilizar en este prototipo se dividen de la siguiente manera:": (
        "Las herramientas seleccionadas para el prototipo se organizan de acuerdo con su función "
        "dentro del pipeline: estimación de pose, seguimiento espacial, clasificación conductual, "
        "postprocesamiento temporal, persistencia de datos e interfaz de usuario."
    ),
    "Técnicas de visión por computadora.": "Estimación de pose y seguimiento por visión por computadora.",
    "Arquitectura de aprendizaje profundo (Deep Learning).": (
        "Arquitecturas de aprendizaje profundo para extracción de keypoints."
    ),
    "Algoritmos de clasificación (Clustering).": (
        "Clasificación supervisada y apoyo no supervisado para análisis conductual."
    ),
    "Herramientas estadísticas.": "Métricas de evaluación, suavizado temporal y validación Leave-One-Out.",
    "Conceptos etológicos para análisis de ansiedad.": (
        "Conceptos etológicos para Grooming, Thigmotaxis y análisis de ansiedad."
    ),
    "Se analizaron las siguientes técnicas de detección de objetos:": (
        "Se analizaron las siguientes técnicas de detección y localización visual:"
    ),
    (
        "Con base en la evaluación de la Tabla 6, se seleccionó la arquitectura YOLOv11 "
        "(redes convolucionales) para la detección y delimitación inicial del espécimen, ya que "
        "su capacidad de extracción espacial en tiempo real y su robustez ante deformaciones "
        "geométricas superan las limitaciones de Haar Cascades y HOG."
    ): (
        "Con base en la evaluación de la Tabla 6, se seleccionó YOLOv11-Pose como arquitectura "
        "de visión por computadora, ya que permite localizar al espécimen y estimar sus keypoints "
        "en un solo flujo de inferencia. Esta decisión evita separar la detección del animal y la "
        "estimación de postura en dos herramientas distintas."
    ),
    (
        "Se optó por utilizar YOLO (You Only Look Once) sobre otras herramientas debido a que su "
        "tiempo de inferencia fue menor que varias implementaciones de Faster R-CNN con solo un "
        "tiempo de 13.5 milisegundos [55]. También suele ser más preciso que Retina Net y "
        "EfficientDet por márgenes de más de 40% [56]. Por otra parte, Adhikari et al sostienen "
        "que YOLO equilibra mejores factores como la precisión, velocidad, robustez y facilidad "
        "de uso [57]."
    ): (
        "En la implementación final se entrenó un modelo YOLO11 Pose con 3,953 imágenes "
        "etiquetadas, alcanzando mAP50 = 0.995 en el conjunto de validación interno. En operación, "
        "el modelo procesa videos de 5 minutos en aproximadamente 3 minutos con GPU dedicada, lo "
        "que lo hace adecuado para el análisis local de sesiones experimentales."
    ),
    "4.2.2 Análisis de técnicas de estimación de pose con aprendizaje profundo": (
        "4.2.2 Análisis comparativo de frameworks de estimación de pose"
    ),
    "Se analizaron las siguientes técnicas de estimación de pose con aprendizaje profundo:": (
        "Se analizaron los siguientes frameworks de estimación de pose animal:"
    ),
    (
        "A pesar de que DLC (DeepLabCut), DPK (DeepPoseKit) y SLEAP poseen características similares "
        "entre sí. DLC ofrece una mejor precisión a diferencia de DPK y SLEAP, dicha precisión de "
        "DLC puede ser comparada con la precisión humana. DLC es más robusto que LEAP y SLEAP para "
        "manejar oclusiones [58]. DLC puede ser entrenado para ser preciso con datasets pequeños "
        "de 100 a 200 imágenes, en comparación con otras herramientas que utilizan mayores cantidades "
        "de datos y no alcanzan una precisión como la de DLC [59]."
    ): (
        "DeepLabCut, SLEAP y DeepPoseKit se conservaron como antecedentes metodológicos por su "
        "relevancia en estimación de pose animal, pero no forman parte del prototipo productivo. "
        "La selección final fue YOLOv11-Pose mediante Ultralytics, debido a su integración directa "
        "con Python, menor tiempo de inferencia local y compatibilidad con el flujo Streamlit del "
        "sistema."
    ),
    "4.2.3 Análisis de técnicas de estimación de pose": (
        "4.2.3 Selección de estimación de pose markerless para el prototipo"
    ),
    "Se analizaron las siguientes técnicas de estimación de pose:": (
        "Se analizaron las siguientes alternativas para capturar postura y trayectoria del espécimen:"
    ),
    (
        "Según Lauer et al, Pose Estimation basada en DLC suele ser precisa sin necesidad de un marcador "
        "además de que es adaptable a varias especies de animales [60]. Mientras tanto, Kinect Skeleton "
        "y Vicon requieren de marcadores y es limitado en entornos naturales [61]."
    ): (
        "Se eligió estimación de pose markerless porque no requiere colocar marcadores físicos en el "
        "animal y, por tanto, no altera la conducta observada en el laberinto. En la versión final, "
        "esta estimación se implementa con YOLOv11-Pose, no con DeepLabCut."
    ),
    "4.2.3.1 Actualización de análisis de técnicas de estimación de pose": (
        "4.2.3.1 Justificación de la migración a YOLOv11-Pose"
    ),
    (
        "El sistema migró de DeepLabCut (DLC) a YOLOv11-Pose debido a que YOLO aplica un método sin mapas "
        "de calor con entrenamiento de extremo a extremo y optimización directa sobre puntos clave, según "
        "Maji et al. [107]. Esta migración solucionó el cuello de botella computacional del hardware local "
        "en los laboratorios de la ENMyH, reduciendo el tiempo de procesamiento de un video de 5 minutos "
        "de 5 horas (con DLC) a solo 3 minutos (con YOLOv11), garantizando una tasa de inferencia superior "
        "a 30 FPS en GPU dedicada y facilitando la distribución autónoma del prototipo en Streamlit."
    ): (
        "La migración a YOLOv11-Pose solucionó el cuello de botella computacional detectado con enfoques "
        "basados en DeepLabCut. Con el modelo actual, un video de 5 minutos se procesa en aproximadamente "
        "3 minutos con GPU dedicada, mientras que las pruebas iniciales con DLC requerían varias horas "
        "por video en el entorno local. Esta decisión también simplificó la integración con el resto del "
        "pipeline, ya que los keypoints se exportan directamente hacia el puente YOLO-SimBA."
    ),
    (
        "De forma conjunta, se optó por utilizar las siguientes arquitecturas:"
    ): (
        "Se analizaron arquitecturas temporales como opciones de apoyo para eventos conductuales complejos:"
    ),
    "RNN (Red Neuronal Recurrente).": "RNN/GRU como referencia para modelar secuencias temporales.",
    "LSTM (Long Short-Term Memory).": "LSTM como mecanismo experimental de rescate temporal.",
    "Transformers.": "Transformers como alternativa futura, no integrada al prototipo productivo.",
    (
        "Según Cole, RNN incorpora cada nueva secuencia y la salida depende de las secuencias ingresadas "
        "con anterioridad, además de que maneja entradas y salidas de longitud variable [67]. Por otra "
        "parte, las LSTM retienen información por largos períodos de tiempos capturando tendencias de "
        "largo plazo [68]. Mientras tanto, los Transformers suelen examinar las secuencias de entrada "
        "simultáneamente en vez de paso a paso [69]."
    ): (
        "En el prototipo final, la clasificación productiva no depende de Transformers ni de una red "
        "recurrente como núcleo principal. El sistema utiliza Random Forest de SimBA para clasificar "
        "Grooming y Thigmotaxis a partir de features espaciotemporales; la capa LSTM se conserva como "
        "apoyo experimental para rescatar eventos de Grooming fragmentados cuando el clasificador base "
        "queda en zona de incertidumbre."
    ),
    (
        "Se optó por utilizar K-Means ya que reduce la complejidad de una imagen ya que se puede utilizar "
        "como herramienta de apoyo para segmentar poblaciones [70]. Por otra parte, Random Forest ayuda "
        "a clasificar el comportamiento animal a partir de datos de postura o movimiento, según Rew et al "
        "[71]. Las Máquinas de Soporte Vectorial (SVM) son una opción robusta para clasificar comportamientos "
        "de animales cuando se cuentan con datos de etiquetado para el entrenamiento, según Mosquera et al [72]."
    ): (
        "Para la clasificación conductual se seleccionó Random Forest mediante SimBA, ya que permite trabajar "
        "con features tabulares derivadas de los keypoints, ofrece interpretabilidad y funcionó de forma "
        "estable en la validación Leave-One-Out. K-Means y SVM se mantuvieron como antecedentes comparativos, "
        "pero no forman parte del flujo productivo. Como apoyo experimental se evaluó B-SOiD, que utiliza "
        "reducción de dimensionalidad y agrupamiento no supervisado para identificar motivos de movimiento."
    ),
    (
        "Se optó por usar Windows como sistema operativo del prototipo debido a su compatibilidad con herramientas "
        "de Deep Learning además de que hay soporte técnico comercial que evita incompatibilidades entre herramientas."
    ): (
        "Se optó por Windows como sistema operativo del prototipo debido a la disponibilidad de equipos del "
        "laboratorio, compatibilidad con GPU NVIDIA/CUDA, Python, Streamlit, OpenCV y Ultralytics. Esta decisión "
        "también facilita la operación por usuarios no especializados en administración de sistemas."
    ),
    (
        "Se optó por utilizar el lenguaje de programación Python debido a su velocidad de desarrollo y rendimiento "
        "en tareas de Machine Learning."
    ): (
        "Se optó por Python debido a su ecosistema para visión por computadora, aprendizaje automático e interfaces "
        "científicas. El prototipo integra Ultralytics/PyTorch para YOLO Pose, SimBA/scikit-learn para Random Forest, "
        "TensorFlow/Keras para pruebas LSTM, OpenCV para video, Streamlit para interfaz y PostgreSQL/SQLAlchemy para "
        "persistencia."
    ),
}


TABLE_UPDATES = {
    9: [
        ["Modelo / Técnica", "Precisión Espacial", "Velocidad de Inferencia", "Carga de Cómputo", "Resultado de Análisis"],
        [
            "YOLOv11-Pose / Redes convolucionales",
            "Excelente (detección del espécimen y keypoints)",
            "Muy alta con GPU dedicada",
            "Baja en inferencia",
            "Seleccionado. Integra localización y estimación de postura en el mismo pipeline.",
        ],
        [
            "Haar Cascades",
            "Baja (sensible a iluminación y postura)",
            "Alta",
            "Muy baja",
            "Descartado. No mantiene precisión ante cambios de orientación y deformación corporal.",
        ],
        [
            "HOG + Linear SVM",
            "Media (útil en formas rígidas)",
            "Media-baja",
            "Media",
            "Descartado. No modela adecuadamente movimientos rápidos ni posturas variables del espécimen.",
        ],
    ],
    10: [
        ["Técnica / Framework", "Precisión de Keypoints", "Integración Local", "Costo Computacional", "Resultado de Análisis"],
        [
            "YOLOv11-Pose (Ultralytics)",
            "Excelente en el dataset local (mAP50 = 0.995)",
            "Alta (Python, OpenCV, Streamlit)",
            "Bajo en inferencia con GPU",
            "Seleccionado como backend productivo de pose.",
        ],
        [
            "DeepLabCut (DLC)",
            "Alta en estimación animal markerless",
            "Media-baja (entornos y dependencias específicos)",
            "Alto en pruebas locales",
            "Usado como antecedente; no integrado al prototipo final por tiempo de procesamiento.",
        ],
        [
            "SLEAP",
            "Alta, especialmente en escenarios multi-animal",
            "Media",
            "Medio-alto",
            "Descartado. El experimento usa un espécimen por video y no requiere su flujo multi-animal.",
        ],
        [
            "DeepPoseKit",
            "Buena",
            "Media-baja",
            "Medio",
            "Descartado por menor mantenimiento y menor alineación con el pipeline actual.",
        ],
    ],
    11: [
        ["Técnica / Dispositivo", "Precisión Espacial", "Requerimiento de Marcadores", "Naturaleza de Captura", "Resultado de Análisis"],
        [
            "YOLOv11-Pose markerless",
            "Excelente para keypoints y centroide",
            "Ninguno",
            "Video estándar del EPM",
            "Seleccionado. Permite registrar conducta natural sin instrumentar al animal.",
        ],
        [
            "Kinect Skeleton Tracking",
            "Baja para roedores",
            "Ninguno",
            "Sensor infrarrojo/profundidad",
            "Descartado. Optimizado para humanos y con resolución insuficiente para roedores.",
        ],
        [
            "Sistemas Vicon / infrarrojo activo",
            "Excelente",
            "Marcadores físicos reflectantes",
            "Cámaras especializadas",
            "Descartado. Instrumentar al animal introduciría variables de estrés y costo adicional.",
        ],
    ],
    12: [
        ["Backend de Pose", "Tasa de Inferencia Local", "Tiempo de Análisis (video 5 min)", "Facilidad de Integración", "Resultado de Análisis"],
        [
            "YOLOv11-Pose (seleccionado)",
            "Muy alta con GPU dedicada",
            "Aproximadamente 3 minutos",
            "Alta (Python, Ultralytics, OpenCV, Streamlit)",
            "Seleccionado. Resuelve el cuello de botella computacional y se integra al puente YOLO-SimBA.",
        ],
        [
            "DeepLabCut (DLC)",
            "Baja en el entorno local evaluado",
            "Varias horas por video en pruebas iniciales",
            "Media-baja por dependencias específicas",
            "Reemplazado. Se conserva como antecedente metodológico, no como backend final.",
        ],
    ],
    14: [
        ["Arquitectura Red", "Propósito en Visión", "Uso en el Prototipo", "Resultado de Análisis"],
        [
            "CNN (Convolutional Neural Network)",
            "Extracción de características espaciales y detección de patrones visuales",
            "Base arquitectónica de YOLOv11-Pose",
            "Seleccionada de forma indirecta mediante YOLO Pose.",
        ],
        [
            "MLP (Multi-Layer Perceptron)",
            "Clasificación de vectores planos",
            "No se usa como arquitectura de visión principal",
            "Descartado para estimación de pose por ignorar estructura espacial de la imagen.",
        ],
    ],
    16: [
        ["Lenguaje", "Ecosistema de IA / ML", "Velocidad de Ejecución", "Velocidad de Desarrollo", "Resultado de Análisis"],
        [
            "Python (seleccionado)",
            "Excelente (Ultralytics/PyTorch, scikit-learn, TensorFlow/Keras, OpenCV, Streamlit)",
            "Media-alta mediante librerías optimizadas",
            "Muy rápida",
            "Seleccionado por integrar visión, clasificación, UI y persistencia en un solo ecosistema.",
        ],
        [
            "C++",
            "Buena para OpenCV nativo",
            "Muy alta",
            "Lenta",
            "Descartado por aumentar tiempo de desarrollo e integración con la UI.",
        ],
        [
            "Java",
            "Limitada para pose animal moderna",
            "Alta",
            "Media",
            "Descartado por menor disponibilidad de herramientas actuales para análisis etológico.",
        ],
    ],
}


def set_tnr_run(run, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def replace_paragraph_text(paragraph, text):
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    run = paragraph.add_run(text)
    set_tnr_run(run)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    set_tnr_run(run)


def resize_table(table, row_count):
    while len(table.rows) < row_count:
        table.add_row()
    while len(table.rows) > row_count:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)


def update_table(table, rows):
    resize_table(table, len(rows))
    for r_idx, row_values in enumerate(rows):
        for c_idx, value in enumerate(row_values):
            set_cell_text(table.cell(r_idx, c_idx), value, bold=(r_idx == 0))


def main():
    doc = Document(DOCX)

    replaced = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in PARAGRAPH_REPLACEMENTS:
            replace_paragraph_text(paragraph, PARAGRAPH_REPLACEMENTS[text])
            replaced += 1

    for table_index, rows in TABLE_UPDATES.items():
        update_table(doc.tables[table_index], rows)

    doc.save(DOCX)
    print(f"Parrafos reemplazados: {replaced}")
    print(f"Tablas actualizadas: {', '.join(str(k) for k in TABLE_UPDATES)}")


if __name__ == "__main__":
    main()
