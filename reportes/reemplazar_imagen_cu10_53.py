from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


DOCX = Path("reportes/reporte_tt2_portocarrero_r_habid/DocumentoTecnicoTT_Habid_V1_copia_codex.docx")
IMG = Path("reportes/figuras/secuencia_53/fig_seq_cu10_bitacora_auditoria.png")


def clear_non_properties(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def main():
    doc = Document(DOCX)
    capture = False
    image_paras = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "5.3 Diagramas de secuencia":
            capture = True
            continue
        if capture and text == "5.4 Diagramas de código":
            break
        if capture and paragraph._p.xpath(".//a:blip"):
            image_paras.append(paragraph)

    if len(image_paras) != 10:
        raise RuntimeError(f"Se esperaban 10 imágenes en 5.3 y se encontraron {len(image_paras)}.")

    target = image_paras[9]
    clear_non_properties(target)
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = target.add_run()
    run.add_picture(str(IMG), width=Inches(6.35))

    doc.save(DOCX)
    print("Imagen CU10 reemplazada en 5.3.")


if __name__ == "__main__":
    main()
